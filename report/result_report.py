
import pandas as pd
from docx import Document
import time
import io
from docx.shared import Inches
import matplotlib.pyplot as plt
import plotly.io as pio
pio.kaleido.scope.default_format = "jpeg" 

class CalibrationReport:
    def __init__(self, auto_calib, calibration_runner):
        self.auto_calib = auto_calib
        self.calibration_runner = calibration_runner
        self.info = []
        self.data_order = {}
        self.data_all = {}
        self.generate_report_data()

    def generate_report_data(self):
        var, standard = self.auto_calib.important_vars
        his = self.calibration_runner.history

        for i in range(len(self.auto_calib.column_num)):
            tar = his[i][1]
            self.info.append(tar)

        for i in range(len(self.info)):
            hf = 0
            hl = 0
            hf_hl = 0
            hf_hl_ti = 0
            r2 = []
            for j in range(len(self.info[i])):
                indi_info = self.info[i][j]
                hl_err = indi_info['洪量相对误差']
                hf_err = indi_info['洪峰相对误差'] 
                ti_err = indi_info['误差h']
                r2_err = indi_info['确定性系数']

                if hl_err <= 0.2:
                    hf += 1
                if hf_err <= 0.2:
                    hl += 1
                if hl_err <= 0.2 and hf_err <= 0.2:
                    hf_hl += 1
                if hl_err <= 0.2 and hf_err <= 0.2 and ti_err <= 2:
                    hf_hl_ti += 1
                r2.append(r2_err)

            hl_err_all = hl / len(self.info[i]) * 100
            hf_err_all = hf / len(self.info[i]) * 100
            hf_hl_err_all = hf_hl / len(self.info[i]) * 100
            hf_hl_ti_err_all = hf_hl_ti / len(self.info[i]) * 100

            r2_mean = sum(r2) / len(r2)
            if r2_mean >= 0.9:
                self.data_order['确定性系数'] = '甲级'
                num_pass = sum(1 for val in r2 if val >= 0.9)
            elif r2_mean >= 0.7 and r2_mean < 0.9:
                self.data_order['确定性系数'] = '乙级'
                num_pass = sum(1 for val in r2 if val >= 0.7)
            elif r2_mean >= 0.5 and r2_mean < 0.7:
                self.data_order['确定性系数'] = '丙级'
                num_pass = sum(1 for val in r2 if val >= 0.5)
            else:
                self.data_order['确定性系数'] = '不合格'
                num_pass = sum(1 for val in r2 if val < 0.3)
            r2per = num_pass / len(r2)

            for err_all, order_type in zip(
                [hl_err_all, hf_err_all, hf_hl_err_all, hf_hl_ti_err_all],
                ['洪量', '洪峰', '洪峰、洪量', '洪峰、洪量、峰现时间']
            ):
                if err_all >= 85:
                    self.data_order[f'{order_type}'] = '甲级'
                elif err_all >= 70:
                    self.data_order[f'{order_type}'] = '乙级'
                elif err_all >= 60:
                    self.data_order[f'{order_type}'] = '丙级'

            section_key = f'断面{i+1}'
            self.data_all[section_key] = [
                {'预报项目': f'洪峰', '总场次': len(self.info[i]), '合格场次': hf, '合格率（%）': hf_err_all, '等级': self.data_order['洪峰']},
                {'预报项目': f'洪量', '总场次': len(self.info[i]), '合格场次': hl, '合格率（%）': hl_err_all, '等级': self.data_order['洪量']},
                {'预报项目': f'洪峰、洪量', '总场次': len(self.info[i]), '合格场次': hf_hl, '合格率（%）': hf_hl_err_all, '等级': self.data_order['洪峰、洪量']},
                {'预报项目': f'洪峰、洪量、峰现时间', '总场次': len(self.info[i]), '合格场次': hf_hl_ti, '合格率（%）': hf_hl_ti_err_all, '等级': self.data_order['洪峰、洪量、峰现时间']},
                {'预报项目': f'确定性系数', '总场次': len(self.info[i]), '合格场次': num_pass, '合格率（%）': r2per, '等级': self.data_order['确定性系数']},
            ]
  

    def add_dataframe_to_docx(self, df, doc):
        table = doc.add_table(rows=df.shape[0] + 1, cols=df.shape[1])
        for j, column in enumerate(df.columns):
            table.cell(0, j).text = str(column)
        for i in range(df.shape[0]):
            for j in range(df.shape[1]):
                table.cell(i + 1, j).text = str(df.iat[i, j])

    def generate_doc(self, doc):
        doc.add_heading("参数率定报告", level=1).alignment = 1 
        doc.add_heading("【全部预报断面的报告】", level=2).alignment = 1

        for i in range(len(self.info)):
            doc.add_heading("断面{}洪水分析".format(i+1), level=3).alignment = 1
            self.add_dataframe_to_docx(pd.DataFrame(self.info[i]), doc)

        doc.add_heading("【整个模型的精度评定报告】", level=2).alignment = 1
        doc.add_paragraph('据《水文情报预报规范》（GB/T22482-2008）要求进行评定，评定结果见下表')
        for i in range(len(self.data_all)):
            doc.add_heading("断面{}洪水分析".format(i+1), level=3).alignment = 1
            self.add_dataframe_to_docx(pd.DataFrame(self.data_all[f'断面{i+1}']), doc)

        doc.add_heading("单个场次的精度图像分析", level=2).alignment = 1
        for i in range(len(self.info)):
            doc.add_heading("断面{}洪水分析".format(i+1), level=3).alignment = 1
            for j in range(len(self.info[i])):
                ID = f'section_{i+1}_{j+1}'
                doc.add_heading(f"第{j+1}场次洪水分析", level=4).alignment = 1
                doc.add_picture(F'./result/{ID}.png', width=Inches(5))  

        doc.add_paragraph(f'\n 报告生成于: {time.strftime("%Y-%m-%d %H:%M:%S")}')
        doc.save('report.docx')


class Cbr(CalibrationReport):
    pass

